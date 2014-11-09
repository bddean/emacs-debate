"""This program converts debate files from the Verbatim template for
ms word to to pseudo-xml format I am using. 

More information on verbatim can be found here: http://paperlessdebate.com/verbatim/SitePages/Home.aspx

TODO: 

TODO: db8 2 docx doesn't work for headings that open at the
previous line. probably change db8 file format. valid xml? maybe
there's an easy way to convert xml formats

http://stackoverflow.com/questions/2405417/automation-how-to-automate-transforming-doc-to-docx

see format macros in verbatim to get a list of debate formats

TODO: smarter cite highlighting

kind of slow
conversion from word should make pre-header newlines headers
db8 format can be simplified a lot
 - header can be annotation at beginning of line
 - if we don't have nexted styles we can just do like <underline>Lorem Ipsum</
   - maybe <UText</  (every style just one char. or maybe 3char code)
   - <u>Text</>
   - <h1/>Line of text
   - <h2/>Line of text
   - <h3/>Line of text
TODO: make sure works for both python2 and python3
TODO: recursive conversion of directories

"""

from __future__ import print_function

import codecs
from docx import Document, text
from docx.oxml.ns import qn
from lxml import etree

print ("running...")

def style(s1, s2):
    return "<" + s2 + ">" + s1 + "</" + s2 + ">"

def docx_to_db8(infile, outfile):
    with open(infile) as f:
        document = Document(f)
    output = "Content-type: text/debate\n"

    for p in document.paragraphs:
        if p.style == "Heading1":
            output += style(p.text, "pocket") + "\n"
        elif p.style == "Heading2":
            output += style(p.text, "hat") + "\n"
        elif p.style == "Heading3":
            output += style(p.text, "block") + "\n"
        elif p.style == "Heading4":
            output += style(p.text, "tag") + "\n"
        elif p.style == "Normal":
            for run in p.runs:
                # Run doesn't provide underline property, so search the xml 
                if run._r.xml.find("highlight") > -1:
                    output += style(run.text, "highlight")
    
                elif run.underline or run.style and run.style.lower() in ('underline',
                                                                        'styleboldunderline',
                                                                        'stylestylebold12pt',
                                                                        'emphasis',
                                                                        'heading4char'):
                    # highlight underlined or bold text in cites,
                    # underline it in card bodies
                    output += style(run.text, ("underline" if card_body else "highlight"))
                else:
                    output += style(run.text, "card")
            output += "\n"
        card_body = p.style == "Normal"
    
    with codecs.open(outfile, "w", "utf-8") as f:
        f.write(output)


# docx_to_db8("/home/ben/Dropbox/Emory Debate Tubs/Affirmatives/Organ Sales/"+
#               "Organ Sales Aff - v1 - Herndon.docx", 
#             "/home/ben/tub/organ_sales_aff.db8")

def highlight_run(run):
    hl = etree.Element(qn("w:highlight"))
    hl.set(qn("w:val"), "yellow")
    rpr = run._r.get_or_add_rPr()
    rpr.append(hl)
    # Highlighted text is underlined as well
    run.style = "Style Bold Underline"

"""
Maybe a state machine isn't the most efficient way to do this
"""
def db8_to_docx(infile, outfile):
    f = codecs.open(infile, encoding="utf-8")
    document = Document("c:/Users/Ben/.emacs.d/Verbatim.docx") # TODO input should be Verbatim template
    # print "\n".join([s.xml for s in document.styles_part.styles._styles_elm])
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    
    translate_style = {
        "pocket": lambda x : setattr(x, "style", "Heading 1"),
        "hat": lambda x : setattr(x, "style", "Heading 2"),
        "block": lambda x : setattr(x, "style", "Heading 3"),
        "tag": lambda x : setattr(x, "style", "Heading 4"),
        "underline": lambda x : setattr(x, "style", "Style Bold Underline"),
        "highlight": lambda x : highlight_run(x),
        # TODO
        "card": lambda x : None
    }
    headings = {"pocket", "hat", "block", "tag"}

    STYLE, WRITE = 0, 1
    mode = WRITE
    current_style  = ""

    # read a byte at a time
    for char in f.read()[len("Content-type: text/debate\n"):]:
        if mode == STYLE:
            # Escape '<'
            if char == "<":
                run.add_text("<")
                mode = WRITE
            elif char == ">":
                mode = WRITE
                # close style => create new run
                if current_style[0] == "/":
                    pass
                # open style => set style
                else:
                    if current_style in headings:
                        paragraph = document.add_paragraph()
                    run = paragraph.add_run()
                    translate_style[current_style](
                        (paragraph if current_style in headings else run))

                current_style=""
            else:
                current_style += char

        elif mode == WRITE:
            if char == "\n":
                continue_run_style = run.style
                # continue_paragraph_style = paragraph.style
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                run.style = continue_run_style
                # paragraph.style = continue_paragraph_style
            else:
                #TODO compare with run.add_text for efficiency
                if char == "<":
                    mode = STYLE
                else:
                    # No control characters or null bytes. TODO: how
                    # do these get in the file in the first place?
                    try:
                        run.add_text(char) # \342 wat
                    except ValueError:
                        print("warning: invalid character", char)
                        pass
                
    f.close()
    document.save(outfile)
# db8_to_docx( // location 1, location 2)

print("done")
